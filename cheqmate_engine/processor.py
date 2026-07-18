import fitz  # PyMuPDF
import pytesseract
from PIL import Image
from docx import Document
import io
import os
import cv2
import numpy as np
import logging

# Configure Logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

class DocumentProcessor:
    def __init__(self):
        # Allow overriding tesseract path via env var if needed
        # On Windows, try to find it, but don't hard crash if Linux
        if os.name == 'nt':
            paths = [
                r'C:\Program Files\Tesseract-OCR\tesseract.exe',
                r'C:\Program Files (x86)\Tesseract-OCR\tesseract.exe',
                os.path.expandvars(r'%LOCALAPPDATA%\Tesseract-OCR\tesseract.exe')
            ]
            for p in paths:
                if os.path.exists(p):
                    pytesseract.pytesseract.tesseract_cmd = p
                    break

    def extract_text(self, file_path: str) -> str:
        """
        Main entry point. Detects file type and routes to appropriate extractor.
        """
        ext = os.path.splitext(file_path)[1].lower()
        
        try:
            if ext == '.pdf':
                return self._process_pdf(file_path)
            elif ext in ['.docx', '.doc']:
                return self._process_docx(file_path)
            elif ext in ['.png', '.jpg', '.jpeg', '.tiff', '.bmp']:
                return self._process_image(file_path)
            elif ext == '.txt':
                with open(file_path, 'r', encoding='utf-8') as f:
                    return f.read()
            else:
                return ""
        except Exception as e:
            logger.error(f"Error processing {file_path}: {e}")
            return ""

    def _process_docx(self, file_path: str) -> str:
        doc = Document(file_path)
        paragraphs_text = []
        
        def process_para(para):
            para_text = []
            for run in para.runs:
                is_invisible = False
                if run.font.color and run.font.color.rgb:
                    rgb = run.font.color.rgb
                    try:
                        if rgb[0] > 245 and rgb[1] > 245 and rgb[2] > 245:
                            is_invisible = True
                    except Exception:
                        pass
                try:
                    if run.font.size and run.font.size.pt < 3.0:
                        is_invisible = True
                except Exception:
                    pass
                
                if not is_invisible:
                    para_text.append(run.text)
            p_str = "".join(para_text)
            if p_str.strip():
                paragraphs_text.append(p_str)

        for para in doc.paragraphs:
            process_para(para)

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        process_para(para)
                        
        return "\n".join(paragraphs_text)

    def _extract_visible_text(self, page) -> str:
        """
        Extracts text from a PyMuPDF page, filtering out white/invisible text and micro-text.
        """
        try:
            blocks = page.get_text("dict")["blocks"]
        except Exception:
            return page.get_text() # Fallback to plain text on any error
            
        page_text_list = []
        for b in blocks:
            if "lines" in b:
                for l in b["lines"]:
                    line_text = []
                    for s in l["spans"]:
                        # Check size (micro-text filter)
                        try:
                            size = s.get("size", 10)
                            if size < 3.0:
                                continue
                        except Exception:
                            pass
                            
                        # Check color (white text filter)
                        try:
                            color = s.get("color", 0)
                            r = (color >> 16) & 255
                            g = (color >> 8) & 255
                            b_val = color & 255
                            
                            if r > 245 and g > 245 and b_val > 245:
                                continue
                        except Exception:
                            pass
                            
                        line_text.append(s["text"])
                    
                    line_str = "".join(line_text)
                    if line_str.strip():
                        page_text_list.append(line_str)
                        
        return "\n".join(page_text_list)

    def _process_image(self, file_path: str) -> str:
        # Load image
        image = cv2.imread(file_path)
        # Preprocess
        processed_img = self._preprocess_image(image)
        # OCR
        text = pytesseract.image_to_string(processed_img)
        return text

    def _process_pdf(self, file_path: str) -> str:
        doc = fitz.open(file_path)
        full_text = []

        for page_num, page in enumerate(doc):
            # 1. Try to extract plain text
            text = self._extract_visible_text(page)
            if text.strip():
                full_text.append(text)
            
            # 2. Extract and OCR all embedded images
            image_list = page.get_images(full=True)
            for img_index, img in enumerate(image_list):
                try:
                    xref = img[0]
                    base_image = doc.extract_image(xref)
                    image_bytes = base_image["image"]
                    
                    # Convert to numpy array for cv2
                    nparr = np.frombuffer(image_bytes, np.uint8)
                    cv_img = cv2.imdecode(nparr, cv2.IMREAD_COLOR)
                    
                    if cv_img is not None:
                        ocr_text = pytesseract.image_to_string(self._preprocess_image(cv_img))
                        if ocr_text.strip():
                            full_text.append(ocr_text)
                except Exception as img_err:
                    logger.warning(f"Failed to extract/OCR embedded image on page {page_num}: {img_err}")
            
            # 3. If there was basically NO text AND NO embedded images, it might be a flat scanned page.
            if len(text.strip()) < 50 and not image_list:
                logger.info(f"Page {page_num} seems entirely scanned with no explicit image objects. Rendering whole page.")
                pix = page.get_pixmap()
                img_data = pix.tobytes("png")
                nparr = np.frombuffer(img_data, np.uint8)
                img = cv2.imdecode(nparr, cv2.IMREAD_COLOR)
                
                try:
                    ocr_text = pytesseract.image_to_string(self._preprocess_image(img))
                    full_text.append(ocr_text)
                except Exception as ocr_err:
                    logger.warning(f"OCR Failed for rendered page {page_num}: {ocr_err}")
                    full_text.append("[OCR Failed - Scanned Image Detected]")
                
        return "\n".join(full_text)

    def _preprocess_image(self, image):
        """
        Applies grayscale, thresholding, and noise removal to improve OCR.
        """
        gray = cv2.cvtColor(image, cv2.COLOR_BGR2GRAY)
        
        # Binary thresholding (Otsu's method)
        _, thresh = cv2.threshold(gray, 0, 255, cv2.THRESH_BINARY + cv2.THRESH_OTSU)
        
        # Denoising (optional, can be slow)
        # denoised = cv2.fastNlMeansDenoising(thresh, None, 10, 7, 21)
        
        return thresh

    def count_images(self, file_path: str) -> int:
        """
        Counts the number of images/screenshots in the document.
        """
        ext = os.path.splitext(file_path)[1].lower()
        if ext == '.pdf':
            try:
                doc = fitz.open(file_path)
                img_count = 0
                for page in doc:
                    img_count += len(page.get_images())
                return img_count
            except Exception as e:
                logger.error(f"Error counting images in PDF: {e}")
                return 0
        elif ext in ['.docx', '.doc']:
            try:
                doc = Document(file_path)
                return len(doc.inline_shapes)
            except Exception as e:
                logger.error(f"Error counting images in docx: {e}")
                return 0
        elif ext in ['.png', '.jpg', '.jpeg', '.tiff', '.bmp']:
            return 1
        return 0

    def extract_text_from_pages(self, file_path: str, start_page: int, end_page: int) -> str:
        """
        Extracts and OCRs text only within a specific 1-indexed page range (inclusive).
        """
        try:
            doc = fitz.open(file_path)
            total_pages = len(doc)
            
            # Clamp page bounds (1-indexed)
            start_idx = max(0, start_page - 1)
            end_idx = min(total_pages - 1, end_page - 1)
            
            if start_idx > end_idx:
                return ""
                
            full_text = []
            for page_num in range(start_idx, end_idx + 1):
                page = doc[page_num]
                # 1. Try to extract plain text
                text = self._extract_visible_text(page)
                if text.strip():
                    full_text.append(text)
                
                # 2. Extract and OCR all embedded images
                image_list = page.get_images(full=True)
                for img_index, img in enumerate(image_list):
                    try:
                        xref = img[0]
                        base_image = doc.extract_image(xref)
                        image_bytes = base_image["image"]
                        
                        nparr = np.frombuffer(image_bytes, np.uint8)
                        cv_img = cv2.imdecode(nparr, cv2.IMREAD_COLOR)
                        
                        if cv_img is not None:
                            try:
                                ocr_text = pytesseract.image_to_string(self._preprocess_image(cv_img))
                                if ocr_text.strip():
                                    full_text.append(ocr_text)
                            except Exception as ocr_err:
                                logger.warning(f"OCR Failed for embedded image on page {page_num}: {ocr_err}")
                    except Exception as img_err:
                        logger.warning(f"Failed to extract/OCR embedded image on page {page_num}: {img_err}")
                
                # 3. Flat scanned page fallback
                if len(text.strip()) < 50 and not image_list:
                    pix = page.get_pixmap()
                    img_data = pix.tobytes("png")
                    nparr = np.frombuffer(img_data, np.uint8)
                    img = cv2.imdecode(nparr, cv2.IMREAD_COLOR)
                    try:
                        ocr_text = pytesseract.image_to_string(self._preprocess_image(img))
                        full_text.append(ocr_text)
                    except Exception as ocr_err:
                        logger.warning(f"OCR Failed for rendered page {page_num}: {ocr_err}")
                        
            return "\n".join(full_text)
        except Exception as e:
            logger.error(f"Error extracting text from page range {start_page}-{end_page}: {e}")
            return ""

    def count_images_from_pages(self, file_path: str, start_page: int, end_page: int) -> int:
        """
        Counts images only within a specific 1-indexed page range (inclusive).
        """
        ext = os.path.splitext(file_path)[1].lower()
        if ext == '.pdf':
            try:
                doc = fitz.open(file_path)
                total_pages = len(doc)
                start_idx = max(0, start_page - 1)
                end_idx = min(total_pages - 1, end_page - 1)
                
                if start_idx > end_idx:
                    return 0
                    
                img_count = 0
                for page_num in range(start_idx, end_idx + 1):
                    img_count += len(doc[page_num].get_images())
                return img_count
            except Exception as e:
                logger.error(f"Error counting images in page range {start_page}-{end_page}: {e}")
                return 0
        return 0

    def extract_images_from_pages(self, file_path: str, start_page: int, end_page: int) -> str:
        """
        Extracts and OCRs all embedded images within a specific 1-indexed page range (inclusive).
        Returns combined OCR text from all images in that range.
        """
        try:
            doc = fitz.open(file_path)
            total_pages = len(doc)
            start_idx = max(0, start_page - 1)
            end_idx = min(total_pages - 1, end_page - 1)

            if start_idx > end_idx:
                return ""

            ocr_texts = []
            for page_num in range(start_idx, end_idx + 1):
                page = doc[page_num]
                image_list = page.get_images(full=True)
                for img in image_list:
                    try:
                        xref = img[0]
                        base_image = doc.extract_image(xref)
                        image_bytes = base_image["image"]
                        nparr = np.frombuffer(image_bytes, np.uint8)
                        cv_img = cv2.imdecode(nparr, cv2.IMREAD_COLOR)
                        if cv_img is not None:
                            ocr_text = pytesseract.image_to_string(self._preprocess_image(cv_img))
                            if ocr_text.strip():
                                ocr_texts.append(ocr_text)
                    except Exception as img_err:
                        logger.warning(f"Failed to OCR image on page {page_num}: {img_err}")

            return "\n".join(ocr_texts)
        except Exception as e:
            logger.error(f"Error extracting images from pages {start_page}-{end_page}: {e}")
            return ""

    def auto_extract_sections(self, file_path: str) -> list:
        """
        Attempts to automatically detect experiment/lab sections in a PDF manual.
        Returns a list of dicts: [{'tag': 'Experiment 1: Name', 'start_page': X, 'end_page': Y}, ...]
        """
        import re
        sections = []
        try:
            doc = fitz.open(file_path)
            total_pages = len(doc)
            
            # 1. Patterns to identify section headings
            # Pattern A: Explicit Section Markers like Experiment/Exp/Lab/Exercise/Task/Topic/Week/Chapter/Section
            major_pattern = re.compile(
                r'^\s*(?:experiment|exp|lab|exercise|task|topic|week|chapter|section)\b\s*[-:#.]?\s*([0-9]+[a-zA-Z0-9.\-]*|[a-zA-Z]+|[ivxldmIVXLDM]+)',
                re.IGNORECASE
            )
            # Pattern B: Numbered Section Headings like "1. Name", "1.1 Aim", "I - Title", etc.
            numbered_pattern = re.compile(
                r'^\s*([0-9]+(?:\.[0-9]+)*|[IVXLDMivxldm]+)\s*[:.-]\s*(.+)$'
            )
            # Pattern C: Numbered Space title like "1 INTRODUCTION" or "II AIM"
            numbered_space_pattern = re.compile(
                r'^\s*([0-9]+(?:\.[0-9]+)*|[IVXLDMivxldm]+)\s+([A-Za-z][A-Za-z0-9 ]{2,50})$'
            )
            
            detected = [] # list of dict: {"page": page_num_1_indexed, "tag": tag_name, "type": "major"|"numbered", "depth": int}
            
            for page_num in range(total_pages):
                page = doc[page_num]
                text = page.get_text()
                
                # Check the first 20 lines of each page to find headings
                lines = [l.strip() for l in text.split('\n') if l.strip()]
                for line in lines[:20]: # Check first 20 lines on the page
                    # Try Major Pattern first
                    match_major = major_pattern.match(line)
                    if match_major:
                        tag = line
                        if len(tag) > 60:
                            tag = tag[:57] + "..."
                        # Avoid duplicate tags for the same page
                        if not any(d["page"] == page_num + 1 for d in detected):
                            detected.append({
                                "page": page_num + 1,
                                "tag": tag,
                                "type": "major",
                                "depth": 0
                            })
                        break # Only one section heading per page start
                    
                    # Try Numbered Pattern
                    match_num = numbered_pattern.match(line)
                    if match_num:
                        num_part = match_num.group(1)
                        depth = num_part.count('.')
                        tag = line
                        if len(tag) > 60:
                            tag = tag[:57] + "..."
                        if not any(d["page"] == page_num + 1 for d in detected):
                            detected.append({
                                "page": page_num + 1,
                                "tag": tag,
                                "type": "numbered",
                                "depth": depth
                            })
                        break
                        
                    # Try Numbered Space Pattern
                    match_num_sp = numbered_space_pattern.match(line)
                    if match_num_sp:
                        num_part = match_num_sp.group(1)
                        depth = num_part.count('.')
                        tag = line
                        if len(tag) > 60:
                            tag = tag[:57] + "..."
                        if not any(d["page"] == page_num + 1 for d in detected):
                            detected.append({
                                "page": page_num + 1,
                                "tag": tag,
                                "type": "numbered",
                                "depth": depth
                            })
                        break
            
            # Filter detected sections:
            # If we have "major" sections, discard all "numbered" sections.
            # If we only have "numbered" sections, filter them to keep only the ones at the minimum depth (highest level).
            has_major = any(d["type"] == "major" for d in detected)
            
            final_detected = []
            if has_major:
                final_detected = [d for d in detected if d["type"] == "major"]
            elif detected:
                min_depth = min(d["depth"] for d in detected)
                final_detected = [d for d in detected if d["depth"] == min_depth]
            
            # If no sections were auto-detected, fallback to splitting by page ranges
            if not final_detected:
                sections.append({
                    "tag": "Full Manual",
                    "start_page": 1,
                    "end_page": total_pages
                })
            else:
                # Sort by page number
                final_detected.sort(key=lambda x: x["page"])
                
                # Build sections with start and end pages
                for i in range(len(final_detected)):
                    start_page = final_detected[i]["page"]
                    tag = final_detected[i]["tag"]
                    
                    if i + 1 < len(final_detected):
                        end_page = final_detected[i+1]["page"] - 1
                        if end_page < start_page:
                            end_page = start_page
                    else:
                        end_page = total_pages
                        
                    sections.append({
                        "tag": tag,
                        "start_page": start_page,
                        "end_page": end_page
                    })
                    
        except Exception as e:
            logger.error(f"Error auto-extracting sections: {e}")
            
        return sections

if __name__ == "__main__":
    # Test
    proc = DocumentProcessor()
    # print(proc.extract_text("test.pdf"))
