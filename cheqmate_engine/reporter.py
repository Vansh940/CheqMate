import fitz  # PyMuPDF
from docx import Document
import os

def append_report_to_pdf(file_path, report_text):
    doc = fitz.open(file_path)
    page = doc.new_page()
    
    # Simple text insertion
    # For a real report, we'd want better formatting
    rect = fitz.Rect(50, 50, 550, 800)
    page.insert_textbox(rect, report_text, fontsize=12, fontname="helv")
    
    # Save to temp and replace? 
    # Fitz supports incremental save, but let's save to a new file then rename
    temp_path = file_path + ".tmp"
    doc.save(temp_path)
    doc.close()
    
    os.replace(temp_path, file_path)

def append_report_to_docx(file_path, report_text):
    try:
        doc = Document(file_path)
        doc.add_page_break()
        doc.add_heading('CheqMate Plagiarism Report', 0)
        doc.add_paragraph(report_text)
        doc.save(file_path)
    except Exception as e:
        print(f"Error appending directly to docx: {e}")
